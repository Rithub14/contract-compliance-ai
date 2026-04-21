import io
import logging
import os
import platform
import uuid

import fitz
import numpy as np
import pypdfium2 as pdfium
from fastapi import APIRouter, HTTPException, UploadFile, File
from PIL import Image

from app.api.job_store import job_store
from app.api.models import UploadResponse
from app.services.custom_rules_service import parse_compliance_doc

logger = logging.getLogger(__name__)

router = APIRouter()

_ALLOWED_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
}
_MAX_BYTES = 20 * 1024 * 1024
_MIN_TEXT_CHARS = 200
_MAX_OCR_PAGES = 20

# OCR_ENGINE=rapidocr forces RapidOCR even on macOS (useful for local testing)
_USE_APPLE_VISION = platform.system() == "Darwin" and os.getenv("OCR_ENGINE", "auto") != "rapidocr"


def _extract_pdf_text_direct(content: bytes) -> str:
    doc = pdfium.PdfDocument(content)
    return "\n\n".join(page.get_textpage().get_text_range() for page in doc)


def _ocr_pdf(content: bytes) -> str:
    doc = fitz.open(stream=content, filetype="pdf")
    pages_text = []
    for i, page in enumerate(doc):
        if i >= _MAX_OCR_PAGES:
            break
        pix = page.get_pixmap(dpi=200)
        if _USE_APPLE_VISION:
            from ocrmac import ocrmac
            img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
            annotations = ocrmac.OCR(img, language_preference=["de-DE", "en-US"]).recognize()
            page_text = "\n".join(item[0] for item in annotations)
        else:
            from rapidocr import RapidOCR
            engine = RapidOCR()
            img_array = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, 3)
            result, _ = engine(img_array)
            page_text = "\n".join(item[1] for item in result) if result else ""
        pages_text.append(page_text)
        logger.debug("ocr page %d: %d chars", i + 1, len(page_text))
    return "\n\n".join(pages_text)


async def _extract_text(content: bytes, content_type: str) -> str:
    if content_type == "application/pdf":
        text = _extract_pdf_text_direct(content)
        logger.info("pdfium extracted %d chars", len(text))
        if len(text.strip()) >= _MIN_TEXT_CHARS:
            return text
        logger.info("text too short — running OCR")
        text = _ocr_pdf(content)
        logger.info("ocr extracted %d chars", len(text))
        return text

    if content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        from docx import Document
        doc = Document(io.BytesIO(content))
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())

    return content.decode(errors="replace")


@router.post("/upload", response_model=UploadResponse)
async def upload_contract(
    file: UploadFile = File(...),
    compliance_file: UploadFile = File(None),
):
    if file.content_type not in _ALLOWED_TYPES:
        raise HTTPException(status_code=400, detail="Only PDF, DOCX, and TXT files are accepted.")

    content = await file.read()
    if len(content) > _MAX_BYTES:
        raise HTTPException(status_code=413, detail="File exceeds 20 MB limit.")

    custom_rules: list[dict] = []
    if compliance_file:
        compliance_content = await compliance_file.read()
        custom_rules = parse_compliance_doc(compliance_content, compliance_file.content_type or "text/plain")

    job_id = str(uuid.uuid4())
    raw_text = await _extract_text(content, file.content_type or "")

    job_store[job_id] = {
        "job_id": job_id,
        "filename": file.filename,
        "raw_text": raw_text,
        "custom_rules": custom_rules,
        "status": "queued",
        "result": None,
    }

    return UploadResponse(
        job_id=job_id,
        status="queued",
        filename=file.filename or "",
        custom_rule_count=len(custom_rules),
    )
